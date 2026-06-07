import streamlit as st
import pandas as pd
from docx import Document
from PIL import Image
import pytesseract
import io

# Configuration de la page
st.set_page_config(page_title="Mon Robot IA Autonome", page_icon="🤖", layout="wide")

st.title("🤖 Assistant Robot IA Multiplateforme")
st.write("Pilotez vos tâches d'état civil depuis votre mobile ou votre ordinateur.")

# --- OPTIONS DANS LA BARRE LATÉRALE ---
st.sidebar.header("⚙️ Menu du Robot")
application_mode = st.sidebar.selectbox(
    "Choisissez une action :",
    ["Saisie d'un Acte Manuscrit", "Générer un fichier Excel", "Créer un document Word"]
)

# --- MODE 1 : LECTURE DU MANUSCRIT ---
if application_mode == "Saisie d'un Acte Manuscrit":
    st.subheader("📸 Numérisation d'un acte d'état civil")
    st.write("Prenez une photo avec votre portable ou importez un scan.")
    
    # Permet d'utiliser l'appareil photo du portable ou de charger un fichier
    fichier_image = st.file_uploader("Importer l'image de l'acte", type=["jpg", "jpeg", "png"])
    
    if fichier_image is not None:
        image = Image.open(fichier_image)
        st.image(image, caption="Acte importé", use_container_width=True)
        
        if st.button("🤖 Lancer la lecture par l'IA"):
            with st.spinner("Lecture du texte manuscrit en cours..."):
                try:
                    # Extraction du texte
                    texte_extrait = pytesseract.image_to_string(image, lang='fra')
                    st.success("Lecture terminée !")
                    
                    # Zone de texte modifiable sur le téléphone
                    texte_final = st.text_area("Texte extrait (vous pouvez le corriger) :", texte_extrait, height=200)
                    
                    # Bouton pour télécharger le résultat en texte simple
                    st.download_button("Télécharger le texte", texte_final, file_name="acte_extrait.txt")
                except Exception as e:
                    st.error(f"Erreur d'OCR : {e}. Assurez-vous que Tesseract est installé sur le serveur.")

# --- MODE 2 : EXCEL ---
elif application_mode == "Générer un fichier Excel":
    st.subheader("📊 Création de Registre Excel")
    
    # Formulaire de saisie rapide adapté au mobile
    nom = st.text_input("Nom :")
    prenom = st.text_input("Prénom :")
    date_naissance = st.text_input("Date de naissance (JJ/MM/AAAA) :")
    
    if st.button("Ajouter au registre"):
        if nom and prenom:
            # Création du fichier Excel en mémoire
            donnees = {"Nom": [nom], "Prénom": [prenom], "Date Naissance": [date_naissance]}
            df = pd.DataFrame(donnees)
            
            buffer = io.BytesIO()
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                df.to_excel(writer, index=False)
            
            st.success("Données enregistrées avec succès dans le robot !")
            st.download_button(
                label="📥 Télécharger le fichier Excel sur votre appareil",
                data=buffer.getvalue(),
                file_name="registre_civil.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        else:
            st.warning("Veuillez remplir au moins le nom et le prénom.")

# --- MODE 3 : WORD ---
elif application_mode == "Créer un document Word":
    st.subheader("📝 Générateur de Document Word")
    titre_doc = st.text_input("Titre du document :", "Acte officiel")
    contenu_doc = st.text_area("Contenu du document :")
    
    if st.button("Générer le Word"):
        doc = Document()
        doc.add_heading(titre_doc, level=1)
        doc.add_paragraph(contenu_doc)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        
        st.download_button(
            label="📥 Télécharger le fichier Word",
            data=buffer.getvalue(),
            file_name="document_robot.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
